import re
from pathlib import Path

TEXT_EXTENSIONS = {
    ".py", ".js", ".ts", ".jsx", ".tsx", ".java", ".go", ".rs", ".c", ".cpp",
    ".h", ".cs", ".rb", ".php", ".swift", ".kt", ".md", ".txt", ".rst",
    ".json", ".yaml", ".yml", ".toml", ".html", ".css", ".sh", ".bash",
}

BINARY_EXTENSIONS = {".pdf", ".docx", ".xlsx"}

SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | BINARY_EXTENSIONS

IGNORED_DIRS = {
    ".git", ".venv", "venv", "node_modules", "__pycache__", ".mypy_cache",
    "dist", "build", ".next", ".nuxt",
}


def can_index(path: Path) -> bool:
    return path.suffix.lower() in SUPPORTED_EXTENSIONS


# ---------------------------------------------------------------------------
# Extractors — turn binary formats into plain text
# ---------------------------------------------------------------------------

def _is_junk_line(line: str) -> bool:
    """Detect lines that are hex dumps (e.g. '0000 90 90 90 90 ...')."""
    stripped = line.strip()
    if not stripped or len(stripped) < 20:
        return False
    # Match classic hex dump: offset followed by repeated 2-char hex bytes
    if re.match(r'^[0-9A-Fa-f]{4}\s+([0-9A-Fa-f]{2}\s+){4,}', stripped):
        return True
    return False


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            # Filter out junk lines (hex dumps, binary data)
            lines = text.splitlines()
            clean_lines = [l for l in lines if not _is_junk_line(l)]
            clean_text = "\n".join(clean_lines).strip()
            if clean_text:
                # Keep page marker as its own paragraph (double newline after)
                pages.append(f"[Pagina {i + 1}]\n\n{clean_text}")
    return "\n\n".join(pages)


def _extract_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # Include tables
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _extract_xlsx(path: Path) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(str(path), read_only=True, data_only=True)
    parts = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        parts.append(f"[Foglio: {sheet}]")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts)


def _extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext == ".xlsx":
        return _extract_xlsx(path)
    return path.read_text(encoding="utf-8", errors="replace")


# ---------------------------------------------------------------------------
# Text normalization
# ---------------------------------------------------------------------------

def _normalize_text(text: str) -> str:
    """Normalize extracted text for better embedding quality.

    - Join wrapped lines within a paragraph (PDFs break lines by layout,
      not by logic, so a single sentence may span 3-4 lines)
    - Dehyphenate words split at end of line (e.g. 'cyberse-\\ncurity')
    - Preserve paragraph boundaries (blank lines)
    - Collapse excessive whitespace
    """
    if not text:
        return ""

    # Split into paragraphs (one or more blank lines = paragraph boundary)
    paragraphs = re.split(r'\n\s*\n+', text)
    normalized: list[str] = []

    for para in paragraphs:
        lines = [l.strip() for l in para.split('\n') if l.strip()]
        if not lines:
            continue

        # Join lines within the paragraph
        result = lines[0]
        for line in lines[1:]:
            if (
                len(result) >= 2
                and result.endswith('-')
                and result[-2].isalpha()
                and line[:1].isalpha()
            ):
                # Hyphenated word split across lines: join without space
                result = result[:-1] + line
            else:
                result += " " + line

        # Collapse runs of whitespace
        result = re.sub(r'\s+', ' ', result).strip()
        if result:
            normalized.append(result)

    return "\n\n".join(normalized)


def _split_sentences(text: str) -> list[str]:
    """Simple sentence splitter for oversized paragraphs."""
    parts = re.split(r'(?<=[.!?])\s+', text)
    return [p.strip() for p in parts if p.strip()]


# ---------------------------------------------------------------------------
# Chunker
# ---------------------------------------------------------------------------

def chunk_file(
    path: Path,
    chunk_words: int = 350,
    overlap_words: int = 50,
) -> list[dict]:
    """Split a file into overlapping chunks along paragraph boundaries.

    Strategy:
    1. Extract + normalize the text
    2. Split into paragraphs
    3. Greedy-pack paragraphs until we hit ``chunk_words``, emit a chunk
    4. Start next chunk with ~``overlap_words`` of overlap from the previous one
    5. If a single paragraph exceeds ``chunk_words``, fall back to sentence split
    6. Each chunk is prefixed with ``[File: <filename>]`` so the embedding
       captures the source context (big win for filename-based queries)
    """
    try:
        text = _extract_text(path).strip()
    except Exception:
        return []

    if not text:
        return []

    text = _normalize_text(text)
    if not text:
        return []

    source = str(path.resolve())
    file_prefix = f"[File: {path.name}]"

    paragraphs = [p for p in text.split("\n\n") if p.strip()]

    chunks: list[dict] = []
    buf: list[str] = []
    buf_words = 0

    def emit_buf() -> None:
        nonlocal buf, buf_words
        if not buf:
            return
        body = "\n\n".join(buf)
        chunks.append({"source": source, "text": f"{file_prefix}\n\n{body}"})

    def take_overlap() -> tuple[list[str], int]:
        over: list[str] = []
        total = 0
        for p in reversed(buf):
            over.insert(0, p)
            total += len(p.split())
            if total >= overlap_words:
                break
        return over, total

    for para in paragraphs:
        pw = len(para.split())

        # Paragraph bigger than one chunk: split on sentences
        if pw > chunk_words:
            if buf:
                emit_buf()
                buf, buf_words = take_overlap()

            sentences = _split_sentences(para)
            sbuf: list[str] = []
            sbuf_words = 0
            for sent in sentences:
                sw = len(sent.split())
                if sbuf_words + sw > chunk_words and sbuf:
                    body = " ".join(sbuf)
                    chunks.append({
                        "source": source,
                        "text": f"{file_prefix}\n\n{body}",
                    })
                    # Sentence-level overlap
                    over: list[str] = []
                    over_w = 0
                    for s in reversed(sbuf):
                        over.insert(0, s)
                        over_w += len(s.split())
                        if over_w >= overlap_words:
                            break
                    sbuf = over
                    sbuf_words = over_w
                sbuf.append(sent)
                sbuf_words += sw

            if sbuf:
                body = " ".join(sbuf)
                chunks.append({
                    "source": source,
                    "text": f"{file_prefix}\n\n{body}",
                })

            buf = []
            buf_words = 0
            continue

        # Normal paragraph: pack into current chunk
        if buf_words + pw > chunk_words and buf:
            emit_buf()
            buf, buf_words = take_overlap()

        buf.append(para)
        buf_words += pw

    if buf:
        emit_buf()

    return chunks


# ---------------------------------------------------------------------------
# Directory walker
# ---------------------------------------------------------------------------

def iter_files(path: Path):
    """Yield all indexable files under path, skipping ignored directories."""
    if path.is_file():
        if can_index(path):
            yield path
    elif path.is_dir():
        for f in path.rglob("*"):
            if not f.is_file():
                continue
            if any(part in IGNORED_DIRS for part in f.parts):
                continue
            if can_index(f):
                yield f
